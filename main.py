# -*- coding: utf-8 -*-
"""
Created on Tue Dec 27 12:37:52 2016

@author: Swad
"""

import sys
from PyQt5.QtWidgets import QTextEdit, QMainWindow, qApp, QAction, QApplication, QPushButton, QLabel,QProgressBar
from PyQt5.QtGui import *
from PyQt5 import  QtCore
import time
from Take_Input import input
from my_preprocess import process,med,dia,medlink,dialink
import os
import urllib2
from docx import Document
from docx.shared import Inches





class GUI(QMainWindow):


    def __init__(self):
        super(GUI, self).__init__()
        self.initUI()
        self.flag = -1

    def initUI(self):
        exitAction = QAction('Exit', self)
        exitAction.setShortcut('ctrl+Q')
        exitAction.setStatusTip('Exit Application')
        exitAction.triggered.connect(qApp.quit)

        self.statusBar()

        menubar = self.menuBar()
        fileMenu = menubar.addMenu('File')
        fileMenu.addAction(exitAction)

        self.line=QTextEdit(self)
        self.line.move(100,100)
        self.line.resize(550,300)

        #self.line1=QTextEdit(self)
        #self.line1.move(100,50)
        #self.line1.resize(550,30)

        self.btn=QPushButton('Process',self)
        self.btn.resize(200,100)
        self.btn.move(680,200)
        self.btn.clicked.connect(self.buttonEvent)

        self.btnonl=QPushButton('Online',self)
        self.btnonl.resize(90,80)
        self.btnonl.move(680,100)
        self.btnonl.setCheckable(False)
        self.btnonl.setStyleSheet("QPushButton { background-color: grey}")
        self.btnonl.clicked.connect(self.onlineEvent)

        self.btnofl = QPushButton('Offline', self)
        self.btnofl.resize(90, 80)
        self.btnofl.move(790, 100)
        self.btnofl.setCheckable(True)
        self.btnofl.setStyleSheet("QPushButton { background-color: red}")
        self.btnofl.clicked.connect(self.offlineEvent)

        self.btnr = QPushButton('Reset', self)
        self.btnr.resize(90, 80)
        self.btnr.move(680, 320)
        self.btnr.clicked.connect(self.resetEvent)

        self.btnp = QPushButton('Print', self)
        self.btnp.resize(90, 80)
        self.btnp.move(790, 320)
        self.btnp.clicked.connect(self.printEvent)


        self.setGeometry(200,100,1000,500)
        self.setWindowTitle('Doctor Machine')
        self.setWindowIcon(QIcon('img.png'))
        
        self.progress = QProgressBar(self)
        self.progress.setGeometry(200, 60, 250, 20)
        self.progress.setRange(0,100)
        self.myLongTask = TaskThread()
        #self.n=QtCore.pyqtSignal(int)
        
        self.myLongTask.notifyProgress.connect(self.onProgress)
        
        
        self.show()
        
        
        
        
    def onProgress(self, j):
            self.progress.setValue(j)
            #self.line.setText("RUN RUN")
            print "in GUI",j
            
             

    def buttonEvent(self):
        self.myLongTask.start()
        
        
            #self.line.setText("RUN RUN")
            
            
        
        self.active=True
        txt=self.line.toPlainText()
        self.sym=txt
        #txt= 'You just entered text:\n' + txt
        
        input(txt)  
        txt1,plt= process()

        tx2=""
        for c in txt1:
            tx2=tx2 + (str(c)+"\n ")
        if self.flag==-1:
           print self.flag
           txt3='\nDisease:\n ' + tx2 #+ "\n\nMedicine and Treatment:\n" + med(txt1) + "\n\nTests and Diagnosis:\n" + dia(txt1)

        else:
           print  self.flag
           txt3="\nDisease: "+ tx2+"\n\n"

           for c in txt1:

                txt3=txt3 +"For "+ c +"\nTreatment and Medicine:\n\n"+medlink(c) + "\n\nTests and Diagnosis:\n\n" + dialink(c) + "\n\n"

        #self.label.setText(txt)

        self.active=False
        time.sleep(.009)
        self.line.setText(txt3)
        plt.show()
        
        #self.line.append("hello")
        
    
    
            
    def resetEvent(self):
        self.line.setText("")
        self.progress.setValue(0)


    def onlineEvent(self):
        self.flag=1
        self.btnonl.setCheckable(True)
        self.btnonl.setStyleSheet("QPushButton { background-color: green}")
        self.btnofl.setCheckable(False)
        self.btnofl.setStyleSheet("QPushButton { background-color: grey }")
        #self.line.setText("You are in online mode")



    def offlineEvent(self):
        self.flag=-1
        self.btnonl.setCheckable(False)
        self.btnonl.setStyleSheet("QPushButton { background-color: grey}")
        self.btnofl.setCheckable(True)
        self.btnofl.setStyleSheet("QPushButton { background-color: red }")
        #self.line.setText("You are in offline mode")




    def printEvent(self):

        txt1 = self.line.toPlainText()

        txt,plt=process()
        plt.savefig("foo.png")
        tx2=""
        for c in txt:
            tx2=tx2 + (str(c)+"\n ")

        #txt1 = self.line.toPlainText()
        #import os
        #all_subdirs = [d for d in os.listdir('.') if os.path.isdir(d)]
        #for file in os.listdir("/home/swad/PycharmProjects/project"):
            #if file.endswith(".png"):
            #    print(file)
        #import glob
        #newest = max(glob.iglob('*.png'), key=os.path.getctime)
        #print newest
        document = Document()
        from docx.shared import Pt

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(7)


        document.add_heading('Disease Details\n\n', level=1)
        p=document.add_paragraph(tx2)
        p.style = document.styles['Normal']
        document.add_picture("foo.png", width=Inches(5))
        document.add_heading(txt1, level=1)
        p2=document.add_paragraph(txt)
        p2.style = document.styles['Normal']
        document.add_page_break()

        document.save('demo.docx')

        self.line.setText("\nDone!")
        
        
        
class TaskThread(QtCore.QThread):
    notifyProgress =QtCore.pyqtSignal(int)
    
    def __init__(self):
        QtCore.QThread.__init__(self)
    
    def run(self):
        self.i=0
        self.s=""
        
        
        while self.i<=100:
              self.notifyProgress.emit(self.i)
              #print "l= ",GUI.l
              #GUI.onProgress(i)
              #GUI.line.setText("in RUNNN")
              #print "in RUN!!!!!!!!!!!!!!!!!!!!  ",self.i
              self.s="RUN RUN RUN"
              #glo=self.i
              
              self.i+=1
              time.sleep(.008)
              

if __name__ == '__main__':
    app=QApplication(sys.argv)
    prog=GUI()
    sys.exit(app.exec_())
